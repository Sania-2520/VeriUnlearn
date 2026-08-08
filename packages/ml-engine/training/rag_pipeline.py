import hashlib
import json
import logging
import os
import re
import time
import uuid
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

try:
    from qdrant_client import QdrantClient
    from qdrant_client.models import (
        Distance,
        FieldCondition,
        Filter,
        MatchValue,
        PointStruct,
        VectorParams,
    )

    QDRANT_AVAILABLE = True
except ImportError:
    QDRANT_AVAILABLE = False

try:
    from sentence_transformers import SentenceTransformer

    ST_AVAILABLE = True
except ImportError:
    ST_AVAILABLE = False

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------


@dataclass
class DocumentChunk:
    chunk_id: str
    document_id: str
    content: str
    metadata: dict
    embedding: Optional[list[float]] = None
    chunk_index: int = 0
    start_char: int = 0
    end_char: int = 0
    token_count: int = 0

    def to_point(self) -> Any:
        if not QDRANT_AVAILABLE:
            raise RuntimeError("qdrant-client is not installed")
        payload = {
            "chunk_id": self.chunk_id,
            "document_id": self.document_id,
            "content": self.content,
            "metadata": self.metadata,
            "chunk_index": self.chunk_index,
            "start_char": self.start_char,
            "end_char": self.end_char,
            "token_count": self.token_count,
        }
        # Stable point ID derived from the chunk ID so re-upserts (e.g. after
        # embedding regeneration) overwrite the same point instead of
        # accumulating duplicates in Qdrant.
        return PointStruct(
            id=f"chunk:{self.chunk_id}",
            vector=self.embedding or [],
            payload=payload,
        )


@dataclass
class Document:
    document_id: str
    filename: str
    content_type: str
    file_size: int
    status: str  # processing / indexed / failed
    chunk_count: int = 0
    total_tokens: int = 0
    metadata: dict = field(default_factory=dict)
    created_at: str = field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )
    indexed_at: Optional[str] = None
    error_message: Optional[str] = None
    file_hash: str = ""

    def to_dict(self) -> dict:
        return {
            "document_id": self.document_id,
            "filename": self.filename,
            "content_type": self.content_type,
            "file_size": self.file_size,
            "status": self.status,
            "chunk_count": self.chunk_count,
            "total_tokens": self.total_tokens,
            "metadata": self.metadata,
            "created_at": self.created_at,
            "indexed_at": self.indexed_at,
            "error_message": self.error_message,
            "file_hash": self.file_hash,
        }

    @classmethod
    def from_dict(cls, data: dict) -> "Document":
        return cls(**{k: v for k, v in data.items() if k in cls.__dataclass_fields__})


@dataclass
class RAGConfig:
    embedding_model: str = "BAAI/bge-large-en-v1.5"
    # ``default_factory`` (not a plain default) so env vars are read at
    # instantiation time, not at import time — this is what lets the Docker
    # deployment override ``QDRANT_URL`` / ``RAG_STORAGE_PATH`` per container.
    qdrant_url: str = field(
        default_factory=lambda: os.getenv("QDRANT_URL", "http://localhost:6333")
    )
    qdrant_collection: str = "documents"
    vector_size: int = 1024
    chunk_size: int = 512
    chunk_overlap: int = 64
    batch_size: int = 32
    max_results: int = 10
    min_score: float = 0.3
    use_hybrid_search: bool = True
    # ``RAG_STORAGE_PATH`` lets the Docker deployment point the pipeline's
    # local metadata store at a volume shared with the backend/worker
    # (``/data/rag``), so document metadata survives restarts in one place.
    storage_path: str = field(
        default_factory=lambda: os.getenv("RAG_STORAGE_PATH", "./rag_storage")
    )


# ---------------------------------------------------------------------------
# Text chunker
# ---------------------------------------------------------------------------


class TextChunker:
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 64):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self._sentence_end_re = re.compile(r"(?<=[.!?])\s+")

    # ------------------------------------------------------------------ #

    def _estimate_tokens(self, text: str) -> int:
        return len(text) // 4

    # ------------------------------------------------------------------ #

    def _split_sentences(self, text: str) -> list[str]:
        return [s for s in self._sentence_end_re.split(text) if s.strip()]

    # ------------------------------------------------------------------ #

    def chunk_text(
        self, text: str, metadata: Optional[dict] = None
    ) -> list[DocumentChunk]:
        metadata = metadata or {}
        sentences = self._split_sentences(text)
        chunks: list[DocumentChunk] = []
        current_sentences: list[str] = []
        current_tokens = 0
        char_offset = 0

        for sentence in sentences:
            sent_tokens = self._estimate_tokens(sentence)
            if current_tokens + sent_tokens > self.chunk_size and current_sentences:
                chunk_content = " ".join(current_sentences).strip()
                start_char = char_offset
                end_char = char_offset + len(chunk_content)
                chunks.append(
                    DocumentChunk(
                        chunk_id=str(uuid.uuid4()),
                        document_id=metadata.get("document_id", ""),
                        content=chunk_content,
                        metadata=dict(metadata),
                        chunk_index=len(chunks),
                        start_char=start_char,
                        end_char=end_char,
                        token_count=self._estimate_tokens(chunk_content),
                    )
                )
                overlap_sentences: list[str] = []
                overlap_tokens = 0
                for prev in reversed(current_sentences):
                    if overlap_tokens + self._estimate_tokens(prev) > self.chunk_overlap:
                        break
                    overlap_sentences.insert(0, prev)
                    overlap_tokens += self._estimate_tokens(prev)
                # Next chunk starts at the end of this chunk (char offsets are
                # sequential, not overlap-relative).
                char_offset = end_char
                current_sentences = overlap_sentences
                current_tokens = overlap_tokens

            current_sentences.append(sentence)
            current_tokens += sent_tokens

        if current_sentences:
            chunk_content = " ".join(current_sentences).strip()
            chunks.append(
                DocumentChunk(
                    chunk_id=str(uuid.uuid4()),
                    document_id=metadata.get("document_id", ""),
                    content=chunk_content,
                    metadata=dict(metadata),
                    chunk_index=len(chunks),
                    start_char=char_offset,
                    end_char=char_offset + len(chunk_content),
                    token_count=self._estimate_tokens(chunk_content),
                )
            )

        logger.info("Chunked text into %d chunks", len(chunks))
        return chunks

    # ------------------------------------------------------------------ #

    def chunk_markdown(
        self, text: str, metadata: Optional[dict] = None
    ) -> list[DocumentChunk]:
        metadata = metadata or {}
        sections = re.split(r"\n(?=#{1,6}\s)", text)
        chunks: list[DocumentChunk] = []
        char_offset = 0

        for section in sections:
            if not section.strip():
                char_offset += len(section)
                continue

            section_tokens = self._estimate_tokens(section)
            if section_tokens <= self.chunk_size:
                chunks.append(
                    DocumentChunk(
                        chunk_id=str(uuid.uuid4()),
                        document_id=metadata.get("document_id", ""),
                        content=section.strip(),
                        metadata=dict(metadata),
                        chunk_index=len(chunks),
                        start_char=char_offset,
                        end_char=char_offset + len(section.strip()),
                        token_count=self._estimate_tokens(section.strip()),
                    )
                )
            else:
                sub_chunks = self.chunk_text(
                    section, metadata
                )
                for sc in sub_chunks:
                    sc.chunk_index = len(chunks)
                    sc.start_char += char_offset
                    sc.end_char += char_offset
                    chunks.append(sc)
            char_offset += len(section)

        if not chunks and text.strip():
            chunks = self.chunk_text(text, metadata)

        logger.info("Markdown chunked into %d chunks", len(chunks))
        return chunks


# ---------------------------------------------------------------------------
# Document processor
# ---------------------------------------------------------------------------


class DocumentProcessor:
    def __init__(self) -> None:
        # Keyed by both MIME type and file extension so callers can pass either
        # (the backend RAG Celery tasks pass extensions such as "pdf"/"docx").
        self._type_map: dict[str, Any] = {
            "application/pdf": self.process_pdf,
            "pdf": self.process_pdf,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self.process_docx,
            "docx": self.process_docx,
            "text/plain": self.process_txt,
            "txt": self.process_txt,
            "text/markdown": self.process_markdown,
            "md": self.process_markdown,
            "markdown": self.process_markdown,
            "text/csv": self.process_csv,
            "text/x-csv": self.process_csv,
            "csv": self.process_csv,
            # Scanned-image OCR (the backend upload endpoint routes image
            # uploads to the ``ocr_process`` Celery task, which lands here).
            "image/png": self.process_image,
            "png": self.process_image,
            "image/jpeg": self.process_image,
            "jpg": self.process_image,
            "jpeg": self.process_image,
            "image/webp": self.process_image,
            "webp": self.process_image,
            "image/tiff": self.process_image,
            "tiff": self.process_image,
            "image/bmp": self.process_image,
            "bmp": self.process_image,
        }

    # ------------------------------------------------------------------ #

    def process_pdf(self, file_path: str) -> str:
        try:
            import pytesseract
            from pdf2image import convert_from_path

            images = convert_from_path(file_path, dpi=300)
            pages = [
                pytesseract.image_to_string(img) for img in images
            ]
            text = "\n\n".join(pages)
            if text.strip():
                return text
        except Exception as exc:
            logger.warning("OCR-based PDF extraction failed (%s), trying fallback", exc)

        try:
            import PyPDF2  # type: ignore

            reader = PyPDF2.PdfReader(file_path)
            pages = [p.extract_text() or "" for p in reader.pages]
            text = "\n\n".join(pages)
            if text.strip():
                return text
        except Exception as exc:
            logger.warning("PyPDF2 fallback failed (%s)", exc)

        try:
            import pytesseract
            from pdf2image import convert_from_path

            images = convert_from_path(file_path, dpi=150)
            text = "\n\n".join(
                pytesseract.image_to_string(img) for img in images
            )
            return text
        except Exception as exc:
            raise RuntimeError(f"Failed to extract text from PDF: {exc}") from exc

    # ------------------------------------------------------------------ #

    def process_image(self, file_path: str) -> str:
        """OCR a scanned image (PNG/JPEG/WebP/TIFF/BMP) into text.

        Uses pytesseract + Pillow. Raises RuntimeError if neither is
        installed or the image contains no extractable text.
        """
        try:
            import pytesseract
            from PIL import Image  # type: ignore
        except ImportError as exc:
            raise RuntimeError(
                "Image OCR requires Pillow and pytesseract — install "
                "ml-engine OCR extras (pytesseract, pdf2image, Pillow)"
            ) from exc

        try:
            with Image.open(file_path) as img:
                text = pytesseract.image_to_string(img)
            if text.strip():
                return text
            raise RuntimeError("OCR returned no text for image")
        except Exception as exc:
            if isinstance(exc, RuntimeError):
                raise
            logger.warning("Image OCR failed (%s), falling back to raw text", exc)
            # Last resort: some "image" files are actually text-based
            # (e.g. mislabeled .txt files); let the txt handler decide.
            return self.process_txt(file_path)

    # ------------------------------------------------------------------ #

    def process_docx(self, file_path: str) -> str:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n\n".join(paragraphs)

    # ------------------------------------------------------------------ #

    def process_txt(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="replace") as fh:
            return fh.read()

    # ------------------------------------------------------------------ #

    def process_markdown(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="replace") as fh:
            return fh.read()

    # ------------------------------------------------------------------ #

    def process_csv(self, file_path: str) -> str:
        import csv

        rows: list[str] = []
        with open(file_path, "r", encoding="utf-8", errors="replace", newline="") as fh:
            reader = csv.reader(fh)
            headers: list[str] = []
            for i, row in enumerate(reader):
                if i == 0:
                    headers = row
                    continue
                row_text = " | ".join(
                    f"{headers[j] if j < len(headers) else f'col{j}'}: {val}"
                    for j, val in enumerate(row)
                )
                rows.append(row_text)
        return "\n\n".join(rows)

    # ------------------------------------------------------------------ #

    def process(self, file_path: str, content_type: str) -> str:
        handler = self._type_map.get(content_type)
        if handler is None:
            if file_path.lower().endswith(".md"):
                return self.process_markdown(file_path)
            if file_path.lower().endswith(".csv"):
                return self.process_csv(file_path)
            return self.process_txt(file_path)
        return handler(file_path)

    # ------------------------------------------------------------------ #

    def get_supported_types(self) -> list[str]:
        return list(self._type_map.keys())


# ---------------------------------------------------------------------------
# Embedding service
# ---------------------------------------------------------------------------


class EmbeddingService:
    def __init__(self, model_name: str = "BAAI/bge-large-en-v1.5"):
        self.model_name = model_name
        self._model: Any = None
        self._dimension: int = 0

    # ------------------------------------------------------------------ #

    def _load_model(self) -> None:
        if self._model is not None:
            return
        if not ST_AVAILABLE:
            logger.warning(
                "sentence-transformers not installed – using random embeddings"
            )
            self._dimension = 1024
            return
        logger.info("Loading embedding model %s", self.model_name)
        self._model = SentenceTransformer(self.model_name)
        # get_embedding_dimension is the current API; fall back to the
        # deprecated alias for older sentence-transformers releases.
        get_dim = getattr(self._model, "get_embedding_dimension", None)
        if get_dim is not None:
            self._dimension = get_dim()
        else:
            self._dimension = self._model.get_sentence_embedding_dimension()
        logger.info("Model loaded – dimension %d", self._dimension)

    # ------------------------------------------------------------------ #

    def embed_texts(
        self, texts: list[str], batch_size: int = 32
    ) -> list[list[float]]:
        self._load_model()
        if self._model is None:
            import random

            return [
                [random.random() for _ in range(self._dimension)] for _ in texts  # nosec B311 - non-crypto dev fallback embeddings
            ]

        all_embeddings: list[list[float]] = []
        for i in range(0, len(texts), batch_size):
            batch = texts[i : i + batch_size]
            emb = self._model.encode(
                batch, normalize_embeddings=True, show_progress_bar=False
            )
            all_embeddings.extend(emb.tolist())
        return all_embeddings

    # ------------------------------------------------------------------ #

    def embed_query(self, query: str) -> list[float]:
        self._load_model()
        if self._model is None:
            import random

            return [random.random() for _ in range(self._dimension)]  # nosec B311 - non-crypto dev fallback embeddings

        prefix = "Represent this sentence for searching relevant passages: "
        emb = self._model.encode(
            prefix + query, normalize_embeddings=True, show_progress_bar=False
        )
        return emb.tolist()

    # ------------------------------------------------------------------ #

    def get_dimension(self) -> int:
        self._load_model()
        return self._dimension

    # ------------------------------------------------------------------ #

    @staticmethod
    def compute_similarity(a: list[float], b: list[float]) -> float:
        import math

        dot = sum(x * y for x, y in zip(a, b))
        norm_a = math.sqrt(sum(x * x for x in a))
        norm_b = math.sqrt(sum(x * x for x in b))
        if norm_a == 0 or norm_b == 0:
            return 0.0
        return dot / (norm_a * norm_b)


# ---------------------------------------------------------------------------
# Vector store  (Qdrant + in-memory fallback)
# ---------------------------------------------------------------------------


class VectorStore:
    def __init__(self, config: RAGConfig):
        self.config = config
        self._client: Any = None
        self._collection_initialized: bool = False
        # In-memory fallback
        self._mem_store: dict[str, dict] = {}

    # ------------------------------------------------------------------ #

    def _init_client(self) -> None:
        if self._client is not None:
            return
        if QDRANT_AVAILABLE:
            try:
                self._client = QdrantClient(
                    url=self.config.qdrant_url, timeout=10, check_version=False
                )
                logger.info("Connected to Qdrant at %s", self.config.qdrant_url)
                return
            except Exception as exc:
                logger.warning(
                    "Cannot connect to Qdrant (%s) – falling back to in-memory store",
                    exc,
                )
        logger.info("Using in-memory vector store")
        self._client = None

    # ------------------------------------------------------------------ #

    def _ensure_collection(self) -> None:
        if self._collection_initialized:
            return
        self._init_client()

        if self._client is not None and QDRANT_AVAILABLE:
            try:
                collections = self._client.get_collections().collections
                names = [c.name for c in collections]
                if self.config.qdrant_collection not in names:
                    self._client.create_collection(
                        collection_name=self.config.qdrant_collection,
                        vectors_config=VectorParams(
                            size=self.config.vector_size,
                            distance=Distance.COSINE,
                        ),
                    )
                    logger.info(
                        "Created Qdrant collection %s",
                        self.config.qdrant_collection,
                    )
            except Exception as exc:
                logger.warning(
                    "Failed to ensure Qdrant collection (%s) – falling back to in-memory store",
                    exc,
                )
                self._client = None
        self._collection_initialized = True

    # ------------------------------------------------------------------ #

    def upsert_chunks(self, chunks: list[DocumentChunk]) -> int:
        self._ensure_collection()

        if self._client is not None and QDRANT_AVAILABLE:
            points = [c.to_point() for c in chunks]
            for i in range(0, len(points), self.config.batch_size):
                batch = points[i : i + self.config.batch_size]
                self._client.upsert(
                    collection_name=self.config.qdrant_collection, points=batch
                )
            return len(chunks)

        # In-memory fallback
        for chunk in chunks:
            self._mem_store[chunk.chunk_id] = {
                "chunk_id": chunk.chunk_id,
                "document_id": chunk.document_id,
                "content": chunk.content,
                "metadata": chunk.metadata,
                "embedding": chunk.embedding or [],
                "chunk_index": chunk.chunk_index,
            }
        return len(chunks)

    # ------------------------------------------------------------------ #

    def search(
        self,
        query_embedding: list[float],
        top_k: int = 5,
        filters: Optional[dict] = None,
    ) -> list[dict]:
        self._ensure_collection()

        if self._client is not None and QDRANT_AVAILABLE:
            qdrant_filter = None
            if filters:
                conditions = []
                for key, value in filters.items():
                    conditions.append(
                        FieldCondition(key=key, match=MatchValue(value=value))
                    )
                qdrant_filter = Filter(must=conditions)

            results = self._client.search(
                collection_name=self.config.qdrant_collection,
                query_vector=query_embedding,
                limit=top_k,
                query_filter=qdrant_filter,
            )
            return [
                {
                    "chunk_id": r.payload.get("chunk_id", ""),
                    "score": r.score,
                    "content": r.payload.get("content", ""),
                    "metadata": r.payload.get("metadata", {}),
                    "document_id": r.payload.get("document_id", ""),
                }
                for r in results
                if r.score >= self.config.min_score
            ]

        # In-memory fallback
        scored: list[dict] = []
        for chunk_id, entry in self._mem_store.items():
            emb = entry.get("embedding", [])
            if not emb:
                continue
            if filters:
                match = all(
                    entry.get("metadata", {}).get(k) == v
                    or entry.get(k) == v
                    for k, v in filters.items()
                )
                if not match:
                    continue
            score = EmbeddingService.compute_similarity(query_embedding, emb)
            if score >= self.config.min_score:
                scored.append(
                    {
                        "chunk_id": entry["chunk_id"],
                        "score": score,
                        "content": entry["content"],
                        "metadata": entry.get("metadata", {}),
                        "document_id": entry.get("document_id", ""),
                    }
                )
        scored.sort(key=lambda x: x["score"], reverse=True)
        return scored[:top_k]

    # ------------------------------------------------------------------ #

    def hybrid_search(
        self, query: str, query_embedding: list[float], top_k: int = 5
    ) -> list[dict]:
        self._ensure_collection()

        keywords = [
            w.lower()
            for w in re.findall(r"\w+", query)
            if len(w) > 2
        ]

        if self._client is not None and QDRANT_AVAILABLE:
            vector_results = self.search(query_embedding, top_k=top_k * 3)
            keyword_filtered = []
            for r in vector_results:
                content_lower = r["content"].lower()
                if any(kw in content_lower for kw in keywords):
                    keyword_filtered.append(r)
            if keyword_filtered:
                return keyword_filtered[:top_k]
            return vector_results[:top_k]

        # In-memory hybrid
        keyword_matches: list[dict] = []
        vector_matches: list[dict] = []
        for chunk_id, entry in self._mem_store.items():
            emb = entry.get("embedding", [])
            content_lower = entry.get("content", "").lower()
            kw_score = sum(1 for kw in keywords if kw in content_lower) / max(
                len(keywords), 1
            )
            vec_score = (
                EmbeddingService.compute_similarity(query_embedding, emb)
                if emb
                else 0.0
            )
            combined = 0.4 * kw_score + 0.6 * vec_score
            if combined >= self.config.min_score:
                item = {
                    "chunk_id": entry["chunk_id"],
                    "score": combined,
                    "content": entry["content"],
                    "metadata": entry.get("metadata", {}),
                    "document_id": entry.get("document_id", ""),
                }
                vector_matches.append(item)
                if kw_score > 0:
                    keyword_matches.append(item)

        vector_matches.sort(key=lambda x: x["score"], reverse=True)
        if keyword_matches:
            keyword_matches.sort(key=lambda x: x["score"], reverse=True)
            return keyword_matches[:top_k]
        return vector_matches[:top_k]

    # ------------------------------------------------------------------ #

    def delete_document(self, document_id: str) -> int:
        self._ensure_collection()

        if self._client is not None and QDRANT_AVAILABLE:
            try:
                results = self._client.scroll(
                    collection_name=self.config.qdrant_collection,
                    scroll_filter=Filter(
                        must=[
                            FieldCondition(
                                key="document_id",
                                match=MatchValue(value=document_id),
                            )
                        ]
                    ),
                    limit=10000,
                )
                points = results[0]
                if points:
                    ids = [p.id for p in points]
                    self._client.delete(
                        collection_name=self.config.qdrant_collection,
                        points_selector=ids,
                    )
                return len(ids)
            except Exception as exc:
                logger.error("Failed to delete document %s: %s", document_id, exc)
                return 0

        to_delete = [
            cid
            for cid, entry in self._mem_store.items()
            if entry.get("document_id") == document_id
        ]
        for cid in to_delete:
            del self._mem_store[cid]
        return len(to_delete)

    # ------------------------------------------------------------------ #

    def get_document_chunks(self, document_id: str) -> list[DocumentChunk]:
        self._ensure_collection()

        if self._client is not None and QDRANT_AVAILABLE:
            try:
                results, _ = self._client.scroll(
                    collection_name=self.config.qdrant_collection,
                    scroll_filter=Filter(
                        must=[
                            FieldCondition(
                                key="document_id",
                                match=MatchValue(value=document_id),
                            )
                        ]
                    ),
                    limit=10000,
                )
                chunks = []
                for pt in results:
                    payload = pt.payload or {}
                    chunks.append(
                        DocumentChunk(
                            chunk_id=payload.get("chunk_id", ""),
                            document_id=payload.get("document_id", ""),
                            content=payload.get("content", ""),
                            metadata=payload.get("metadata", {}),
                            chunk_index=payload.get("chunk_index", 0),
                            start_char=payload.get("start_char", 0),
                            end_char=payload.get("end_char", 0),
                            token_count=payload.get("token_count", 0),
                        )
                    )
                return chunks
            except Exception as exc:
                logger.error("Failed to get chunks for %s: %s", document_id, exc)
                return []

        return [
            DocumentChunk(
                chunk_id=e["chunk_id"],
                document_id=e.get("document_id", ""),
                content=e["content"],
                metadata=e.get("metadata", {}),
                chunk_index=e.get("chunk_index", 0),
            )
            for e in self._mem_store.values()
            if e.get("document_id") == document_id
        ]

    # ------------------------------------------------------------------ #

    def count_chunks(self) -> int:
        self._ensure_collection()
        if self._client is not None and QDRANT_AVAILABLE:
            try:
                info = self._client.get_collection(self.config.qdrant_collection)
                return info.points_count or 0
            except Exception:
                logger.warning("Qdrant collection count failed")
                return 0
        return len(self._mem_store)

    # ------------------------------------------------------------------ #

    def upsert_vector(
        self,
        collection: str,
        point_id: str,
        vector: list[float],
        payload: Optional[dict] = None,
    ) -> None:
        """Upsert a raw vector into an arbitrary collection.

        Used by the memory pipeline for semantic memory vectors that are not
        tied to a RAG document. Falls back to an in-memory store when Qdrant
        is unavailable.
        """
        self._ensure_collection()
        payload = dict(payload or {})
        payload["point_id"] = point_id
        if self._client is not None and QDRANT_AVAILABLE:
            try:
                collections = self._client.get_collections().collections
                names = [c.name for c in collections]
                if collection not in names:
                    self._client.create_collection(
                        collection_name=collection,
                        vectors_config=VectorParams(
                            size=len(vector) or self.config.vector_size,
                            distance=Distance.COSINE,
                        ),
                    )
                self._client.upsert(
                    collection_name=collection,
                    points=[
                        PointStruct(id=point_id, vector=vector, payload=payload)
                    ],
                )
                return
            except Exception as exc:
                logger.warning(
                    "Qdrant vector upsert failed (%s) – falling back to in-memory", exc
                )
        self._mem_store[f"{collection}:{point_id}"] = {
            "collection": collection,
            "point_id": point_id,
            "embedding": vector,
            **payload,
        }

    # ------------------------------------------------------------------ #

    def delete_by_filter(self, collection: str, filter_: Optional[dict] = None) -> int:
        """Delete vectors matching an exact-match payload filter.

        Safety: an empty filter is a no-op (refuses to wipe a whole collection).
        In-memory matching covers both raw vectors (stored under
        ``{collection}:{point_id}`` keys) and document chunks (stored under
        bare ``chunk_id`` keys), so document cleanup by ``document_id`` works
        regardless of which store is active.
        """
        self._ensure_collection()
        filter_ = filter_ or {}
        if not filter_:
            logger.warning("delete_by_filter called with an empty filter — refusing to delete")
            return 0

        if self._client is not None and QDRANT_AVAILABLE:
            try:
                results, _ = self._client.scroll(
                    collection_name=collection,
                    scroll_filter=Filter(
                        must=[
                            FieldCondition(key=k, match=MatchValue(value=v))
                            for k, v in filter_.items()
                        ]
                    ),
                    limit=10000,
                )
                if results:
                    self._client.delete(
                        collection_name=collection,
                        points_selector=[p.id for p in results],
                    )
                return len(results)
            except Exception as exc:
                logger.error("Qdrant delete_by_filter failed: %s", exc)
                return 0

        prefix = f"{collection}:"
        to_delete = []
        for key, entry in self._mem_store.items():
            # Raw vectors are keyed ``{collection}:{point_id}``; document
            # chunks are keyed by bare ``chunk_id`` (match those too when the
            # caller filters by a payload field such as ``document_id``).
            is_raw_vector = key.startswith(prefix)
            is_chunk = "chunk_id" in entry and entry.get("chunk_id")
            if not (is_raw_vector or is_chunk):
                continue
            if all(entry.get(k) == v for k, v in filter_.items()):
                to_delete.append(key)
        for key in to_delete:
            del self._mem_store[key]
        return len(to_delete)

    # ------------------------------------------------------------------ #

    def get_collection_stats(self) -> dict:
        self._ensure_collection()
        if self._client is not None and QDRANT_AVAILABLE:
            try:
                info = self._client.get_collection(self.config.qdrant_collection)
                return {
                    "collection": self.config.qdrant_collection,
                    "vectors_count": info.vectors_count or 0,
                    "points_count": info.points_count or 0,
                    "status": str(info.status),
                    "backend": "qdrant",
                }
            except Exception as exc:
                return {"error": str(exc), "backend": "qdrant"}
        return {
            "collection": "in-memory",
            "points_count": len(self._mem_store),
            "backend": "in-memory",
        }


# ---------------------------------------------------------------------------
# RAG Pipeline
# ---------------------------------------------------------------------------


class RAGPipeline:
    def __init__(self, config: Optional[RAGConfig] = None):
        self.config = config or RAGConfig()
        self.chunker = TextChunker(
            chunk_size=self.config.chunk_size,
            chunk_overlap=self.config.chunk_overlap,
        )
        self.processor = DocumentProcessor()
        self.embedding_service = EmbeddingService(
            model_name=self.config.embedding_model
        )
        self.vector_store = VectorStore(self.config)
        self._documents: dict[str, Document] = {}

        os.makedirs(self.config.storage_path, exist_ok=True)
        self._load_documents()

    # ------------------------------------------------------------------ #

    def ingest_document(
        self,
        file_path: str,
        filename: str,
        content_type: str,
        metadata: Optional[dict] = None,
    ) -> Document:
        metadata = metadata or {}

        file_size = os.path.getsize(file_path)
        file_hash = self._hash_file(file_path)

        for doc in self._documents.values():
            if doc.file_hash == file_hash:
                logger.info(
                    "Document %s already indexed (hash match)", filename
                )
                return doc

        document_id = str(uuid.uuid4())
        doc = Document(
            document_id=document_id,
            filename=filename,
            content_type=content_type,
            file_size=file_size,
            status="processing",
            metadata=metadata,
            file_hash=file_hash,
        )
        self._documents[document_id] = doc
        self._save_document_metadata(doc)

        try:
            text = self.processor.process(file_path, content_type)
            if not text.strip():
                raise RuntimeError("Extracted text is empty")

            doc_meta = {**metadata, "document_id": document_id, "filename": filename}

            is_markdown = content_type == "text/markdown" or filename.lower().endswith(
                ".md"
            )
            if is_markdown:
                chunks = self.chunker.chunk_markdown(text, doc_meta)
            else:
                chunks = self.chunker.chunk_text(text, doc_meta)

            if not chunks:
                raise RuntimeError("No chunks produced from document")

            texts = [c.content for c in chunks]
            embeddings = self.embedding_service.embed_texts(
                texts, batch_size=self.config.batch_size
            )
            for chunk, embedding in zip(chunks, embeddings):
                chunk.embedding = embedding

            self.vector_store.upsert_chunks(chunks)

            total_tokens = sum(c.token_count for c in chunks)
            doc.status = "indexed"
            doc.chunk_count = len(chunks)
            doc.total_tokens = total_tokens
            doc.indexed_at = datetime.now(timezone.utc).isoformat()
            self._save_document_metadata(doc)

            logger.info(
                "Ingested %s → %d chunks (%d tokens)",
                filename,
                len(chunks),
                total_tokens,
            )
            return doc

        except Exception as exc:
            doc.status = "failed"
            doc.error_message = str(exc)
            self._save_document_metadata(doc)
            logger.error("Failed to ingest %s: %s", filename, exc)
            raise

    # ------------------------------------------------------------------ #

    def ingest_text(
        self,
        text: str,
        source_name: str,
        metadata: Optional[dict] = None,
    ) -> Document:
        metadata = metadata or {}

        text_hash = hashlib.sha256(text.encode("utf-8")).hexdigest()

        for doc in self._documents.values():
            if doc.file_hash == text_hash:
                return doc

        document_id = str(uuid.uuid4())
        doc = Document(
            document_id=document_id,
            filename=source_name,
            content_type="text/plain",
            file_size=len(text.encode("utf-8")),
            status="processing",
            metadata=metadata,
            file_hash=text_hash,
        )
        self._documents[document_id] = doc
        self._save_document_metadata(doc)

        try:
            doc_meta = {**metadata, "document_id": document_id, "filename": source_name}
            chunks = self.chunker.chunk_text(text, doc_meta)
            if not chunks:
                raise RuntimeError("No chunks produced from text")

            texts = [c.content for c in chunks]
            embeddings = self.embedding_service.embed_texts(
                texts, batch_size=self.config.batch_size
            )
            for chunk, embedding in zip(chunks, embeddings):
                chunk.embedding = embedding

            self.vector_store.upsert_chunks(chunks)

            total_tokens = sum(c.token_count for c in chunks)
            doc.status = "indexed"
            doc.chunk_count = len(chunks)
            doc.total_tokens = total_tokens
            doc.indexed_at = datetime.now(timezone.utc).isoformat()
            self._save_document_metadata(doc)
            return doc

        except Exception as exc:
            doc.status = "failed"
            doc.error_message = str(exc)
            self._save_document_metadata(doc)
            raise

    # ------------------------------------------------------------------ #

    def search(
        self,
        query: str,
        top_k: int = 5,
        filters: Optional[dict] = None,
    ) -> list[dict]:
        query_embedding = self.embedding_service.embed_query(query)

        if self.config.use_hybrid_search:
            results = self.vector_store.hybrid_search(
                query, query_embedding, top_k=top_k
            )
        else:
            results = self.vector_store.search(
                query_embedding, top_k=top_k, filters=filters
            )

        return results

    # ------------------------------------------------------------------ #

    def build_context(self, query: str, top_k: int = 5) -> str:
        results = self.search(query, top_k=top_k)
        if not results:
            return "No relevant context found."

        context_parts: list[str] = []
        for i, r in enumerate(results, 1):
            doc_id = r.get("document_id", "unknown")
            score = r.get("score", 0)
            content = r.get("content", "")
            context_parts.append(
                f"[Source {i} | document_id={doc_id} | relevance={score:.3f}]\n{content}"
            )
        return "\n\n---\n\n".join(context_parts)

    # ------------------------------------------------------------------ #

    def build_prompt(
        self,
        query: str,
        context: str,
        system_prompt: Optional[str] = None,
    ) -> str:
        if system_prompt is None:
            system_prompt = (
                "You are a helpful assistant. Answer the user's question based on "
                "the provided context. If the context does not contain enough "
                "information, say so honestly."
            )

        return (
            f"{system_prompt}\n\n"
            f"## Context\n\n{context}\n\n"
            f"## Question\n\n{query}\n\n"
            f"## Answer\n"
        )

    # ------------------------------------------------------------------ #

    def delete_document(self, document_id: str) -> bool:
        if document_id not in self._documents:
            return False
        self.vector_store.delete_document(document_id)
        doc = self._documents.pop(document_id)
        meta_path = os.path.join(
            self.config.storage_path, f"{document_id}.json"
        )
        if os.path.exists(meta_path):
            os.remove(meta_path)
        logger.info("Deleted document %s", doc.filename)
        return True

    # ------------------------------------------------------------------ #

    def get_document(self, document_id: str) -> Optional[Document]:
        return self._documents.get(document_id)

    # ------------------------------------------------------------------ #

    def list_documents(
        self, page: int = 1, page_size: int = 25
    ) -> tuple[list[Document], int]:
        all_docs = list(self._documents.values())
        all_docs.sort(key=lambda d: d.created_at, reverse=True)
        total = len(all_docs)
        start = (page - 1) * page_size
        end = start + page_size
        return all_docs[start:end], total

    # ------------------------------------------------------------------ #

    def reindex_document(self, document_id: str) -> Document:
        doc = self._documents.get(document_id)
        if doc is None:
            raise KeyError(f"Document {document_id} not found")

        self.vector_store.delete_document(document_id)
        meta_path = os.path.join(
            self.config.storage_path, f"{document_id}.json"
        )
        file_path = meta_path

        stored_path = doc.metadata.get("file_path")
        if stored_path and os.path.exists(stored_path):
            file_path = stored_path
        elif not os.path.exists(meta_path):
            raise FileNotFoundError(
                f"Original file for document {document_id} not found"
            )

        if stored_path and os.path.exists(stored_path):
            return self.ingest_document(
                stored_path,
                doc.filename,
                doc.content_type,
                metadata=doc.metadata,
            )

        raise FileNotFoundError(
            f"Original file for document {document_id} no longer available on disk"
        )

    # ------------------------------------------------------------------ #

    def get_stats(self) -> dict:
        docs = list(self._documents.values())
        indexed = sum(1 for d in docs if d.status == "indexed")
        failed = sum(1 for d in docs if d.status == "failed")
        processing = sum(1 for d in docs if d.status == "processing")
        total_chunks = sum(d.chunk_count for d in docs)
        total_tokens = sum(d.total_tokens for d in docs)

        return {
            "total_documents": len(docs),
            "indexed_documents": indexed,
            "failed_documents": failed,
            "processing_documents": processing,
            "total_chunks": total_chunks,
            "total_tokens": total_tokens,
            "vector_store": self.vector_store.get_collection_stats(),
            "embedding_model": self.config.embedding_model,
            "embedding_dimension": self.embedding_service.get_dimension(),
            "config": {
                "chunk_size": self.config.chunk_size,
                "chunk_overlap": self.config.chunk_overlap,
                "min_score": self.config.min_score,
                "use_hybrid_search": self.config.use_hybrid_search,
            },
        }

    # ------------------------------------------------------------------ #

    def process_document(
        self,
        document_id: str,
        filename: str,
        file_type: str,
        storage_path: str = "",
        text: str = "",
        metadata: Optional[dict] = None,
    ) -> Document:
        """Ingest a document from a storage path (or inline text) under a
        caller-supplied ``document_id`` (Celery task path).

        The document is parsed with the configured :class:`DocumentProcessor`
        (OCR for PDFs, DOCX/CSV/TXT/MD extraction), chunked, embedded and
        indexed into the vector store — the same pipeline used by the
        synchronous upload path.
        """
        metadata = dict(metadata or {})
        metadata["document_id"] = document_id
        metadata["filename"] = filename

        # Re-ingest under a fixed id so the caller (backend DB row) stays in
        # sync with the index: drop any prior record with the same id first.
        prior = self._documents.get(document_id)
        if prior is not None:
            self.delete_document(document_id)

        file_size = 0
        file_hash = ""
        if text:
            file_hash = hashlib.sha256(text.encode("utf-8")).hexdigest()
            file_size = len(text.encode("utf-8"))
        elif storage_path and os.path.exists(storage_path):
            file_size = os.path.getsize(storage_path)
            file_hash = self._hash_file(storage_path)
        else:
            raise FileNotFoundError(
                f"storage_path '{storage_path}' does not exist and no inline text provided"
            )

        doc = Document(
            document_id=document_id,
            filename=filename,
            content_type=file_type,
            file_size=file_size,
            status="processing",
            metadata=metadata,
            file_hash=file_hash,
        )
        self._documents[document_id] = doc
        self._save_document_metadata(doc)

        try:
            if text:
                extracted = text
            else:
                extracted = self.processor.process(storage_path, file_type)
            if not extracted.strip():
                raise RuntimeError("Extracted text is empty")

            # For PDFs the OCR/PDF extraction path joins pages with blank lines;
            # record an honest page count for the caller (Celery worker stores
            # it in the DB row's page_count column).
            if file_type in ("pdf", "application/pdf"):
                page_count = max(1, len([p for p in extracted.split("\n\n") if p.strip()]))
            else:
                page_count = 1
            metadata["page_count"] = page_count

            doc_meta = dict(metadata)
            is_markdown = (
                file_type == "text/markdown"
                or filename.lower().endswith(".md")
            )
            if is_markdown:
                chunks = self.chunker.chunk_markdown(extracted, doc_meta)
            else:
                chunks = self.chunker.chunk_text(extracted, doc_meta)
            if not chunks:
                raise RuntimeError("No chunks produced from document")

            embeddings = self.embedding_service.embed_texts(
                [c.content for c in chunks], batch_size=self.config.batch_size
            )
            for chunk, embedding in zip(chunks, embeddings):
                chunk.embedding = embedding

            self.vector_store.upsert_chunks(chunks)

            doc.status = "indexed"
            doc.chunk_count = len(chunks)
            doc.total_tokens = sum(c.token_count for c in chunks)
            doc.indexed_at = datetime.now(timezone.utc).isoformat()
            self._save_document_metadata(doc)
            logger.info(
                "Processed document %s → %d chunks (%d tokens)",
                filename, len(chunks), doc.total_tokens,
            )
            return doc
        except Exception as exc:
            doc.status = "failed"
            doc.error_message = str(exc)
            self._save_document_metadata(doc)
            logger.error("Failed to process %s: %s", filename, exc)
            raise

    # ------------------------------------------------------------------ #

    def regenerate_embeddings(self, document_id: str) -> int:
        """Re-embed every chunk of an indexed document and refresh the store.

        Returns the number of chunks successfully re-embedded.
        """
        doc = self._documents.get(document_id)
        if doc is None:
            return 0
        chunks = self.vector_store.get_document_chunks(document_id)
        if not chunks:
            return 0

        texts = [c.content for c in chunks]
        embeddings = self.embedding_service.embed_texts(
            texts, batch_size=self.config.batch_size
        )
        refreshed: list[DocumentChunk] = []
        for chunk, embedding in zip(chunks, embeddings):
            chunk.embedding = embedding
            refreshed.append(chunk)
        self.vector_store.upsert_chunks(refreshed)
        logger.info("Regenerated embeddings for %s (%d chunks)", document_id, len(refreshed))
        return len(refreshed)

    # ------------------------------------------------------------------ #

    def ocr_process(
        self,
        document_id: str,
        storage_path: str = "",
        file_type: str = "pdf",
        metadata: Optional[dict] = None,
    ) -> Document:
        """Extract text from a scanned/PDF document via OCR and index it.

        Uses the OCR-first extraction path of :class:`DocumentProcessor`
        (pytesseract + pdf2image with a PyPDF2 fallback).
        """
        if not storage_path or not os.path.exists(storage_path):
            raise FileNotFoundError(
                f"OCR storage_path '{storage_path}' does not exist"
            )
        return self.process_document(
            document_id=document_id,
            filename=os.path.basename(storage_path),
            file_type=file_type,
            storage_path=storage_path,
            metadata=metadata,
        )

    # ------------------------------------------------------------------ #

    def upsert_vector(
        self,
        collection: str,
        point_id: str,
        vector: list[float],
        payload: Optional[dict] = None,
    ) -> None:
        """Upsert a raw vector (memory pipeline) into the vector store."""
        self.vector_store.upsert_vector(collection, point_id, vector, payload)

    def delete_vectors(self, collection: str, filter_: Optional[dict] = None) -> int:
        """Delete vectors matching a payload filter (memory pipeline)."""
        return self.vector_store.delete_by_filter(collection, filter_)

    # ------------------------------------------------------------------ #

    def _save_document_metadata(self, doc: Document) -> None:
        path = os.path.join(
            self.config.storage_path, f"{doc.document_id}.json"
        )
        with open(path, "w", encoding="utf-8") as fh:
            json.dump(doc.to_dict(), fh, indent=2, ensure_ascii=False)

    # ------------------------------------------------------------------ #

    def _load_documents(self) -> None:
        storage = Path(self.config.storage_path)
        if not storage.exists():
            return
        for meta_file in storage.glob("*.json"):
            try:
                with open(meta_file, "r", encoding="utf-8") as fh:
                    data = json.load(fh)
                doc = Document.from_dict(data)
                self._documents[doc.document_id] = doc
            except Exception as exc:
                logger.warning("Failed to load metadata from %s: %s", meta_file, exc)
        logger.info("Loaded %d document metadata records", len(self._documents))

    # ------------------------------------------------------------------ #

    @staticmethod
    def _hash_file(file_path: str) -> str:
        sha = hashlib.sha256()
        with open(file_path, "rb") as fh:
            for block in iter(lambda: fh.read(8192), b""):
                sha.update(block)
        return sha.hexdigest()
